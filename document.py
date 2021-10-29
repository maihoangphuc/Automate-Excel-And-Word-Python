#import library
from docx import Document

#define doc
doc = Document()

#add heading with level 0
doc.add_heading('Thư mời họp phụ huynh', 0)

#add paragraph 1
p = doc.add_paragraph('Học sinh: ')
#add run
p.add_run('Nguyễn Văn Minh').bold = True
p.add_run(' - Lớp 12').italic = True

#add paragraph 2
p = doc.add_paragraph('Thời Gian: ')
#add run
p.add_run('9AM ngày 1/1/2021').bold = True

#add paragraph 3
p = doc.add_paragraph('Địa điểm: ')
#add run
p.add_run('Phòng 12A1 Trường THPT NVL').bold = True

#add heading with level 1
doc.add_heading('Nội dung', 1)

#add paragraph 4
p = doc.add_paragraph('Tổng kế năm học 2021 vừa qua' + 
'và nhận xét quá trình học của từng học sinh:')

#add paragraph 5
p = doc.add_paragraph('Lộ trình học năm 2022', style='List Bullet')
p = doc.add_paragraph('Trao thưởng những học sinh xuất sắc', style='List Bullet')
p = doc.add_paragraph('Giải thích các khoản thu', style='List Bullet')

#add heading with level 1
doc.add_heading('Các khoản thu', 1)

#add paragraph 6
p = doc.add_paragraph('1. Học phí: 6.000.000 VND', style='List Number')
p = doc.add_paragraph('1. Đồng phục: 3.000.000 VND', style='List Number')
p = doc.add_paragraph('1. Tham quan: 2.000.000 VND', style='List Number')

#add heading with level 2
doc.add_heading('Tổng các khoản thu: 11.000.000 VND', 2)

#save file docx
doc.save('./docx/thumoi.docx')



